"""Generate PEERPoint Admin How-To as .docx and .pdf into apps/pwa/public/docs/."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "public" / "docs"
DOCX_PATH = OUT_DIR / "PEERPoint-Admin-How-To.docx"
PDF_PATH = OUT_DIR / "PEERPoint-Admin-How-To.pdf"

TITLE = "PEERPoint — Admin How-To"

SECTIONS: list[tuple[str, list[str]]] = [
    (
        "Sign in (Admin)",
        [
            "Open https://admin.mypeerpoint.com.",
            "Enter your Admin username and password, then click Sign in.",
            "The seed account is username Admin. Change the password under Change password after first login.",
            "Your session lasts about 12 hours, or until you sign out or close the tab.",
        ],
    ),
    (
        "Invite a Peer Support Member",
        [
            "Any Admin can invite Staff or Admin access. Admins automatically have staff queue rights.",
            "Under Invite Peer Support Member, enter: First name, Last name, Bureau, Job title, Email, and Access (Staff or Admin).",
            "Click Send invite. If Resend is configured they get an email; always keep the invite link as a backup.",
            "The invitee finishes registration on /setup: username, password, current shift, cell/home/work phones, personal and work emails.",
            "Pending invites can be Resent or Revoked. Staff sign in at https://mypeerpoint.com/staff.",
        ],
    ),
    (
        "Manage accounts",
        [
            "Disable / Enable blocks or restores sign-in (except the seed Admin account).",
            "Make Admin / Make Staff changes access level (except the seed Admin account).",
            "Each user updates their own password under Change password after login.",
        ],
    ),
    (
        "Run the request queue",
        [
            "Update On duty (comma-separated names) → Save on-duty list.",
            "Set Peer display name for assignments (your name or initials).",
            "Click Refresh to load new help requests.",
            "For an open request: review phone, email, and notes → Assign + generate room code.",
            "Contact the member and give them the room code. The app does not send the code automatically.",
            "Join Peer chat and/or Peer voice with the same room code (on the member site, or share the code with an on-duty peer).",
            "When finished → Close the request → Sign out.",
        ],
    ),
    (
        "What members see",
        [
            "On https://mypeerpoint.com, members can use Request Help, Self Help, Peer chat, and Peer voice.",
            "Request Help collects phone, email, and optional notes for follow-up.",
            "Peer chat / Peer voice require a room code you share after assigning a request.",
            "PEERPoint is peer support, not clinical care or emergency services. Chat and voice are not recorded by the agency in this app.",
            "Crisis options stay visible for members: 988, 911, and the peer support line.",
        ],
    ),
    (
        "Who signs in where",
        [
            "Admin → https://admin.mypeerpoint.com → Username + password.",
            "Staff → https://mypeerpoint.com/staff → Username + password.",
            "Invitee → link from email (/setup?token=…) → completes registration.",
            "Staff username login is rejected on the production Admin site.",
        ],
    ),
    (
        "Troubleshooting",
        [
            "Prefer https://admin.mypeerpoint.com for Admin sign-in.",
            "Staff cannot sign in on Admin site → They must use https://mypeerpoint.com/staff.",
            "Invite email not received → Copy the invite link from the Admin site; check RESEND_API_KEY and INVITE_FROM_EMAIL.",
            "Invite expired → Resend or create a new invite (links expire in 7 days).",
            "Queue empty but a member says they submitted → Refresh; offline/local saves on the member’s device never appear in the queue.",
            "Member never got a room code → Codes are not emailed by the app — you must share them.",
        ],
    ),
    (
        "Admin checklist",
        [
            "Open https://admin.mypeerpoint.com and sign in with your Admin username.",
            "Change the seed Admin password after first login.",
            "Invite Staff/Admin with profile fields + email.",
            "Confirm invitees complete /setup.",
            "Disable accounts as needed.",
            "Keep the on-duty list current; assign and close requests.",
        ],
    ),
]

FOOTER = (
    "Emergencies: 911 · Crisis: 988 · Peer line: 801-548-8002 · "
    "Email: slcosopeersupport@saltlakecounty.gov · "
    "IT: PEERPOINT_KV required; optional RESEND_API_KEY + INVITE_FROM_EMAIL"
)

INTRO = (
    "This guide is for program administrators who invite Peer Support Members and oversee "
    "the help-request queue on PEERPoint."
)

LINKS = [
    ("Admin site", "https://admin.mypeerpoint.com"),
    ("Member & staff site", "https://mypeerpoint.com"),
    ("Staff sign-in", "https://mypeerpoint.com/staff"),
    ("Invite setup", "https://mypeerpoint.com/setup"),
]


def build_docx() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    title = doc.add_heading(TITLE, level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x2F)

    intro = doc.add_paragraph(INTRO)
    intro.paragraph_format.space_after = Pt(10)

    for label, url in LINKS:
        p = doc.add_paragraph()
        run = p.add_run(f"{label}: ")
        run.bold = True
        p.add_run(url)
        p.paragraph_format.space_after = Pt(2)

    doc.add_paragraph()

    for heading, bullets in SECTIONS:
        doc.add_heading(heading, level=1)
        for item in bullets:
            bp = doc.add_paragraph(item, style="List Bullet")
            bp.paragraph_format.space_after = Pt(4)
            bp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    foot = doc.add_paragraph(FOOTER)
    foot.paragraph_format.space_before = Pt(16)
    for run in foot.runs:
        run.italic = True
        run.font.size = Pt(10)

    doc.save(DOCX_PATH)


def build_pdf() -> None:
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "TitleCustom",
        parent=styles["Title"],
        fontSize=18,
        textColor=colors.HexColor("#1b3a2f"),
        spaceAfter=12,
        alignment=TA_LEFT,
    )
    h_style = ParagraphStyle(
        "HCustom",
        parent=styles["Heading2"],
        fontSize=13,
        textColor=colors.HexColor("#1b3a2f"),
        spaceBefore=14,
        spaceAfter=6,
    )
    body = ParagraphStyle(
        "BodyCustom",
        parent=styles["Normal"],
        fontSize=10.5,
        leading=14,
        spaceAfter=4,
    )
    bullet = ParagraphStyle(
        "BulletCustom",
        parent=body,
        leftIndent=14,
        bulletIndent=0,
        spaceAfter=3,
    )
    muted = ParagraphStyle(
        "Muted",
        parent=body,
        fontSize=9.5,
        textColor=colors.HexColor("#44554c"),
        spaceBefore=12,
    )

    story: list = []
    story.append(Paragraph(TITLE, title_style))
    story.append(Paragraph(INTRO, body))
    story.append(Spacer(1, 6))

    link_data = [[Paragraph(f"<b>{label}</b>", body), Paragraph(url, body)] for label, url in LINKS]
    table = Table(link_data, colWidths=[1.6 * inch, 5.2 * inch])
    table.setStyle(
        TableStyle(
            [
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 0),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 1),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 1),
            ]
        )
    )
    story.append(table)
    story.append(Spacer(1, 8))

    for heading, items in SECTIONS:
        story.append(Paragraph(heading, h_style))
        for item in items:
            story.append(Paragraph(f"• {item}", bullet))

    story.append(Paragraph(FOOTER, muted))

    doc = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=letter,
        leftMargin=inch,
        rightMargin=inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
        title=TITLE,
        author="PEERPoint",
    )
    doc.build(story)


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    build_docx()
    build_pdf()
    print(f"Wrote {DOCX_PATH}")
    print(f"Wrote {PDF_PATH}")


if __name__ == "__main__":
    main()
