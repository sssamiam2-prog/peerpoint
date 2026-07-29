"""Generate PEERPoint Admin + Staff How-To docs (.docx and .pdf) into apps/pwa/public/docs/."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "public" / "docs"

FOOTER = (
    "Emergencies: 911 · Crisis: 988 · Peer line: 801-548-8002 · "
    "Email: slcosopeersupport@saltlakecounty.gov · "
    "IT: PEERPOINT_KV required; optional RESEND_API_KEY + INVITE_FROM_EMAIL"
)

ADMIN = {
    "title": "PEERPoint — Admin How-To",
    "intro": (
        "This guide is for program administrators who invite Peer Support Members, manage On Call, "
        "oversee the request queue, and review reports on PEERPoint."
    ),
    "links": [
        ("Admin site", "https://admin.mypeerpoint.com"),
        ("Member & staff site", "https://mypeerpoint.com"),
        ("Staff sign-in", "https://mypeerpoint.com/staff"),
        ("Invite setup", "https://mypeerpoint.com/setup"),
    ],
    "docx": OUT_DIR / "PEERPoint-Admin-How-To.docx",
    "pdf": OUT_DIR / "PEERPoint-Admin-How-To.pdf",
    "sections": [
        (
            "Sign in (Admin)",
            [
                "Open https://admin.mypeerpoint.com.",
                "Enter your Admin username and password, then click Sign in.",
                "The seed account is username Admin. Change the password under Account after first login.",
                "Your session lasts about 12 hours, or until you sign out or close the tab.",
            ],
        ),
        (
            "Invite a Peer Support Member",
            [
                "Any Admin can invite Staff or Admin access. Admins automatically have staff queue rights.",
                "Under Members → Invite Peer Support Member, enter: First name, Last name, Bureau, Job title, Email, and Access (Staff or Admin).",
                "Click Send invite. Invitees receive an email with Finish registration and How-To download buttons.",
                "Always keep the invite link shown in the Admin site as a backup if email is delayed.",
                "The invitee finishes registration on /setup: username, password, sex (for matching), phones, and emails.",
                "Pending invites can be Resent or Revoked. Staff sign in at https://mypeerpoint.com/staff.",
            ],
        ),
        (
            "Manage accounts",
            [
                "Disable / Enable blocks or restores sign-in (except the seed Admin account).",
                "Make Admin / Make Staff changes access level (except the seed Admin account).",
                "Set Male / Female on each account so immediate-contact matching works.",
                "Each user updates their own password under Account after login.",
            ],
        ),
        (
            "On Call schedule",
            [
                "Open the On Call tab. Choose a day and start/end times.",
                "Admins can schedule any Peer Support Member from the roster; Staff can only add themselves.",
                "Before saving, the person being scheduled must acknowledge they are expected to be available during those times.",
                "Members requesting immediate chat/voice are matched to who is On Call now (Male/Female preference). Names are not shown to members.",
            ],
        ),
        (
            "Run the request queue",
            [
                "Set Peer display name for assignments (your name or initials).",
                "Click Refresh to load new help requests (form, face-to-face, or immediate contact).",
                "For an open request: review contact info and notes → Assign + generate room code.",
                "Share the room code with the member. The app does not email the code automatically (except immediate-contact alerts to staff).",
                "Join Peer chat and/or Peer voice with the same room code.",
                "Add staff notes on the request as needed. If you are assigned, log actual time spent in minutes.",
                "When finished → Close the request → Sign out.",
            ],
        ),
        (
            "Content: Self Help, videos, and Resource Gallery",
            [
                "Open the Content tab on the Admin site.",
                "Edit Self Help articles (title, category, body, related link). Paste a YouTube, Vimeo, or .mp4 link to embed a video for members.",
                "Publish or reset to built-in articles when needed.",
                "Add Resource Gallery items by uploading a file (max 8 MB) or pasting an external link. Members open these under Resources on the member site.",
                "Delete outdated gallery files from the same Content tab.",
            ],
        ),
        (
            "Reports",
            [
                "Open the Reports tab on the Admin site.",
                "Review totals: open/assigned/closed requests and total minutes logged.",
                "See time spent by Peer Support Member, request notes, and On Call history (including availability acknowledgments).",
                "Use Refresh on the Reports tab after staff log notes or time.",
            ],
        ),
        (
            "What members see",
            [
                "On https://mypeerpoint.com, members unlock the site with a Site use code, then use Request Help, Self Help, Peer chat, and Peer voice.",
                "Request Help includes employment attestation, bureau, civilian/sworn, and contact preferences.",
                "Immediate contact matches an On Call peer by sex preference and opens chat or voice with a room code.",
                "PEERPoint is peer support, not clinical care or emergency services.",
                "Crisis options stay visible: 988, 911, and the peer support line.",
            ],
        ),
        (
            "Troubleshooting",
            [
                "Prefer https://admin.mypeerpoint.com for Admin sign-in.",
                "Staff cannot sign in on the Admin site → They must use https://mypeerpoint.com/staff.",
                "Invite email not received → Copy the invite link from the Admin site; check RESEND_API_KEY and INVITE_FROM_EMAIL.",
                "Invite expired → Resend or create a new invite (links expire in 7 days).",
                "Queue empty but a member says they submitted → Refresh; offline/local saves never appear in the queue.",
                "Member never got a room code → Codes are not emailed by the app — you must share them (except staff alerts for immediate contact).",
            ],
        ),
        (
            "Admin checklist",
            [
                "Open https://admin.mypeerpoint.com and sign in with your Admin username.",
                "Change the seed Admin password after first login.",
                "Invite Staff/Admin; confirm they complete /setup and receive the How-To.",
                "Keep On Call coverage current with acknowledged availability.",
                "Assign/close requests; review Reports for notes and time spent.",
            ],
        ),
    ],
}

STAFF = {
    "title": "PEERPoint — Staff How-To",
    "intro": (
        "This guide is for Peer Support Members (Staff) who take On Call shifts, respond to help requests, "
        "add notes, and log time spent on PEERPoint."
    ),
    "links": [
        ("Staff sign-in", "https://mypeerpoint.com/staff"),
        ("Member site", "https://mypeerpoint.com"),
        ("Invite setup", "https://mypeerpoint.com/setup"),
    ],
    "docx": OUT_DIR / "PEERPoint-Staff-How-To.docx",
    "pdf": OUT_DIR / "PEERPoint-Staff-How-To.pdf",
    "sections": [
        (
            "Finish registration",
            [
                "Open the Finish registration button from your invite email (or the backup link from Admin).",
                "Choose a username and password. Add phones, emails, current shift, and sex (Male/Female) for matching.",
                "After setup, sign in at https://mypeerpoint.com/staff. Do not use the Admin site unless you are an Admin.",
            ],
        ),
        (
            "Sign in",
            [
                "Open https://mypeerpoint.com/staff.",
                "Enter your username and password.",
                "Your session lasts about 12 hours, or until you sign out or close the tab.",
                "Change your password anytime under the Account tab.",
            ],
        ),
        (
            "Add yourself to On Call",
            [
                "Open the On Call tab.",
                "Your name is selected by default. Choose the day and start/end times you can cover.",
                "Check the acknowledgment that you are expected to be available during those times, then add the block.",
                "You can only schedule yourself. Admins can schedule other Peer Support Members if needed.",
                "When you are On Call and a member requests immediate contact matching your sex preference, you may get an alert email with a Join button and room code.",
            ],
        ),
        (
            "Respond to requests",
            [
                "Open the Requests tab and click Refresh.",
                "Set Peer display name for assignments (your name or initials).",
                "For an open request: Assign + generate room code, then contact the member and share the code.",
                "Use Open Peer chat or Open Peer voice with that room code.",
                "Room codes expire after 24 hours with no chat/voice use; re-assign to issue a new code.",
                "Close the request when finished.",
            ],
        ),
        (
            "Notes and time spent",
            [
                "On any request, add a staff note to record important follow-up details. Notes are visible to Staff and Admins.",
                "If a request is assigned to you, log actual time spent in minutes (optional short note).",
                "Only the assigned Peer Support Member (or an Admin) can log time on that request.",
                "Admins can see notes and time in the Reports tab.",
            ],
        ),
        (
            "What members experience",
            [
                "Members unlock https://mypeerpoint.com with a Site use code (not an employee ID).",
                "They can request help (including face-to-face), use self-help resources, or request immediate chat/voice.",
                "PEERPoint is peer support, not clinical care or emergency services. Point members to 988 / 911 when needed.",
            ],
        ),
        (
            "Troubleshooting",
            [
                "Cannot sign in on admin.mypeerpoint.com → Use https://mypeerpoint.com/staff.",
                "Invite expired → Ask an Admin to Resend the invite (links expire in 7 days).",
                "No On Call match for a member → Confirm your sex is set and you have an active On Call block.",
                "Cannot log time → Assign the request to yourself first (or ask Admin).",
            ],
        ),
        (
            "Staff checklist",
            [
                "Complete /setup from your invite email and download this How-To.",
                "Sign in at https://mypeerpoint.com/staff.",
                "Add yourself to On Call for days/times you can cover (with acknowledgment).",
                "Refresh Requests; assign, chat/voice, add notes, log time, then close.",
                "Sign out when finished.",
            ],
        ),
    ],
}


def build_docx(spec: dict) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    title = doc.add_heading(spec["title"], level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x2F)

    intro = doc.add_paragraph(spec["intro"])
    intro.paragraph_format.space_after = Pt(10)

    for label, url in spec["links"]:
        p = doc.add_paragraph()
        run = p.add_run(f"{label}: ")
        run.bold = True
        p.add_run(url)
        p.paragraph_format.space_after = Pt(2)

    doc.add_paragraph()

    for heading, bullets in spec["sections"]:
        doc.add_heading(heading, level=1)
        for item in bullets:
            bp = doc.add_paragraph(item, style="List Bullet")
            bp.paragraph_format.space_after = Pt(4)
            bp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    foot = doc.add_paragraph(FOOTER)
    foot.paragraph_format.space_before = Pt(16)
    for run in foot.runs:
        run.italic = True
        run.font.size = Pt(10)

    doc.save(spec["docx"])


def build_pdf(spec: dict) -> None:
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "TitleCustom",
        parent=styles["Title"],
        fontSize=18,
        textColor=colors.HexColor("#1b3a2f"),
        spaceAfter=12,
        alignment=TA_LEFT,
    )
    h_style = ParagraphStyle(
        "HCustom",
        parent=styles["Heading2"],
        fontSize=13,
        textColor=colors.HexColor("#1b3a2f"),
        spaceBefore=14,
        spaceAfter=6,
    )
    body = ParagraphStyle(
        "BodyCustom",
        parent=styles["Normal"],
        fontSize=10.5,
        leading=14,
        spaceAfter=4,
    )
    bullet = ParagraphStyle(
        "BulletCustom",
        parent=body,
        leftIndent=14,
        bulletIndent=0,
        spaceAfter=3,
    )
    muted = ParagraphStyle(
        "Muted",
        parent=body,
        fontSize=9.5,
        textColor=colors.HexColor("#44554c"),
        spaceBefore=12,
    )

    story: list = []
    story.append(Paragraph(spec["title"], title_style))
    story.append(Paragraph(spec["intro"], body))
    story.append(Spacer(1, 6))

    link_data = [[Paragraph(f"<b>{label}</b>", body), Paragraph(url, body)] for label, url in spec["links"]]
    table = Table(link_data, colWidths=[1.6 * inch, 5.2 * inch])
    table.setStyle(
        TableStyle(
            [
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 0),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 1),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 1),
            ]
        )
    )
    story.append(table)
    story.append(Spacer(1, 8))

    for heading, items in spec["sections"]:
        story.append(Paragraph(heading, h_style))
        for item in items:
            story.append(Paragraph(f"• {item}", bullet))

    story.append(Paragraph(FOOTER, muted))

    doc = SimpleDocTemplate(
        str(spec["pdf"]),
        pagesize=letter,
        leftMargin=inch,
        rightMargin=inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
        title=spec["title"],
        author="PEERPoint",
    )
    doc.build(story)


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    for spec in (ADMIN, STAFF):
        build_docx(spec)
        build_pdf(spec)
        print(f"Wrote {spec['docx']}")
        print(f"Wrote {spec['pdf']}")


if __name__ == "__main__":
    main()
